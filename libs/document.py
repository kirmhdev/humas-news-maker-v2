import datetime, math, os

# from libs import scrape
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from libs.scrape import scrape_image_to_bytes

# hti = Html2Image()
document = Document()


def create_pages(data, format, headers):
    for i, d in enumerate(data):
        title = document.add_paragraph()
        title_format = title.paragraph_format
        title_format.space_after = Pt(0)
        title_font = title.add_run(d.title).font
        title_font.name = format["titleFont"]
        title_font.size = Pt(format["titleSize"])
        title_font.bold = format["titleBold"]

        try:
            image = scrape_image_to_bytes(d.image, headers)
            document.add_picture(image, height=Inches(format["imageHeight"]))
        except Exception as e:
            print(e)

        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for p in d.paragraphs:
            paragraph = document.add_paragraph()
            paragraph_font = paragraph.add_run(p).font
            paragraph_font.name = format["paragraphFont"]
            paragraph_font.size = Pt(format["paragraphSize"])
        if i + 1 < len(data):
            document.add_page_break()


# def create_classements(classement_list):
#     for classement in classement_list:
#         scrape.get_classement(classement)
#         document.add_picture(".cache/screenshot.png", width=Inches(4.3))


def create_document(data, format, headers):
    create_pages(data, format, headers=headers)

    section = document.sections[0]
    section.page_width = Inches(format["pageWidth"])
    section.page_height = Inches(format["pageHeight"])
    section.top_margin = Inches(format["pageMtop"])
    section.bottom_margin = Inches(format["pageMbot"])
    section.left_margin = Inches(format["pageMlef"])
    section.right_margin = Inches(format["pageMrig"])

    document.add_page_break()

    # if use_classements:
    #     create_classements(classement_list)

    datetime_now = datetime.datetime.now()
    first_weekday = datetime.datetime(
        year=datetime_now.year, month=datetime_now.month, day=1
    ).weekday()

    offset_date = datetime_now.day + first_weekday - 1
    week = math.floor(offset_date / 7)
    month, year = datetime_now.month, datetime_now.year
    if week == 0:
        week = 4
        month -= 1
        if month == 0:
            month = 12
            year -= 1

    formatted_date = datetime.datetime(month=month, year=year, day=1)

    edition = f"{formatted_date.strftime("%B %Y")} ({str(week)})"
    if format["useFooter"]:
        section.footer.add_paragraph(f"Santri Update {edition}")

    out_dir = f"OUT/{formatted_date.strftime("%Y/%B")}"
    os.makedirs(out_dir, exist_ok=True)

    document.save(f"{out_dir}/sans {edition}.docx")
