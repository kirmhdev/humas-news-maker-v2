import datetime, math, os, json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from core.scrape import get_classements, scrape_image_to_bytes

# hti = Html2Image()
document = Document()


def get_out_filename(root_out_dir):
    datetime_now = datetime.datetime.now()

    current_directory = os.getcwd()
    out_dir = f"{root_out_dir}/{datetime_now.strftime("%Y/%B")}"
    final_directory = os.path.join(current_directory, out_dir)
    if not os.path.exists(final_directory):
        os.makedirs(final_directory)

    out_filename = os.path.join(
        final_directory, f"sans {datetime_now.strftime("%d %A")}"
    )
    return out_filename


def create_pages(data, format, headers):
    for i, d in enumerate(data):
        title = document.add_paragraph()
        title_format = title.paragraph_format
        title_format.space_after = Pt(0)
        title_font = title.add_run(d["title"]).font
        title_font.name = format["titleFont"]
        title_font.size = Pt(format["titleSize"])
        title_font.bold = format["titleBold"]

        try:
            image = scrape_image_to_bytes(d["image"], headers)
            document.add_picture(image, height=Inches(format["imageHeight"]))
        except Exception as e:
            print(e)

        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for p in d["paragraphs"]:
            paragraph = document.add_paragraph()
            paragraph_font = paragraph.add_run(p).font
            paragraph_font.name = format["paragraphFont"]
            paragraph_font.size = Pt(format["paragraphSize"])
        if i + 1 < len(data):
            document.add_page_break()


def create_document(data, format, headers, root_out_dir, classement_sources):
    create_pages(data, format, headers=headers)

    section = document.sections[0]
    section.page_width = Inches(format["pageWidth"])
    section.page_height = Inches(format["pageHeight"])
    section.top_margin = Inches(format["pageMtop"])
    section.bottom_margin = Inches(format["pageMbot"])
    section.left_margin = Inches(format["pageMlef"])
    section.right_margin = Inches(format["pageMrig"])

    document.add_page_break()

    for classement_source in classement_sources:
        get_classements(classement_source)
        document.add_picture(".cache/screenshot.png", width=Inches(4.3))

    out_filename = get_out_filename(root_out_dir)

    document.save(f"{out_filename}.docx")


def save_news_to_json(data, root_out_dir):
    out_filename = get_out_filename(root_out_dir)

    try:
        with open(f"{out_filename}.json", "w") as file:
            json.dump(data, file)
    except:
        print("Error when saving news")
